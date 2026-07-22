from __future__ import annotations

from html.parser import HTMLParser
import re
from datetime import date, datetime
from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO
from typing import Any, Iterable


VALID_BUSINESS_TYPES = {"agreement", "income_certificate"}
VALID_BLOCK_TYPES = {"header", "title", "head", "paragraph", "body", "line", "sign", "footer"}
VALID_SOURCE_TYPES = {"employee_field", "computed", "manual", "fixed", "system"}

COMPANY_PHONE = "0755-86685111"
COMPANY_ADDRESS = "广东省深圳市南山区科苑路15号科兴科学园A3栋16楼"

_VAR_PATTERN = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")


def date_text(value: Any) -> str:
    d = _coerce_date(value)
    if not d:
        return ""
    return f"{d.year}年{d.month}月{d.day}日"


def month_text(value: Any) -> str:
    d = _coerce_date(value)
    if not d:
        return ""
    return f"{d.year}年{d.month}月"


def money_text(value: Any) -> str:
    amount = Decimal(str(value or 0)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return f"{amount:,.2f}"


def wan_text(value: Any) -> str:
    amount = Decimal(str(value or 0)) / Decimal("10000")
    return f"{amount.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP):,.2f}"


def _coerce_date(value: Any) -> date | None:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    if value in (None, ""):
        return None
    text = str(value).strip().replace("/", "-")
    for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(text[:19] if " " in fmt else text[:10], fmt).date()
        except ValueError:
            continue
    return None


def _installment_sentence(installments: Iterable[dict[str, Any]]) -> str:
    parts: list[str] = []
    for item in installments:
        d = _coerce_date(item.get("pay_date"))
        if not d:
            continue
        parts.append(f"{date_text(d)}前支付{money_text(item.get('amount'))}元")
    return "；".join(parts) + ("。" if parts else "")


def _period_text(start: Any, end: Any) -> str:
    start_text = date_text(start)
    end_text = date_text(end)
    if start_text and end_text:
        return f"{start_text}至{end_text}"
    if start_text:
        return f"{start_text}至今"
    return "至今"


def enrich_values(business_type: str, values: dict[str, Any]) -> dict[str, Any]:
    merged = dict(values)
    if business_type == "agreement":
        merged.setdefault("dissolve_date_text", date_text(merged.get("dissolve_date")))
        merged.setdefault("last_work_date_text", date_text(merged.get("last_work_date")))
        merged.setdefault("salary_until_text", date_text(merged.get("salary_until")))
        merged.setdefault("social_security_month", month_text(merged.get("dissolve_date")))
        merged.setdefault("base_amount_text", money_text(merged.get("base_amount")))
        merged.setdefault("total_amount_text", money_text(merged.get("total_amount")))
        merged.setdefault("installments_text", _installment_sentence(merged.get("installments") or []))
    elif business_type == "income_certificate":
        merged.setdefault("period_text", _period_text(merged.get("hire_date"), merged.get("leave_date")))
        merged.setdefault("leave_date_text", date_text(merged.get("leave_date")) or "至今")
        merged.setdefault("issue_date_text", date_text(merged.get("issue_date")))
        merged.setdefault("basic_salary_text", money_text(merged.get("basic_salary")))
        merged.setdefault("target_bonus_text", money_text(merged.get("target_bonus")))
        merged.setdefault("annual_package_text", money_text(merged.get("annual_package")))
        merged.setdefault("annual_package_wan", wan_text(merged.get("annual_package")))
        merged.setdefault("company_phone", COMPANY_PHONE)
        merged.setdefault("company_address", COMPANY_ADDRESS)
    return merged


def render_template_blocks(
    blocks: Iterable[Any],
    values: dict[str, Any],
    variables: Iterable[Any] = (),
    business_type: str = "",
) -> list[tuple[str, str]]:
    merged = _values_with_defaults(values, variables)
    if business_type:
        merged = enrich_values(business_type, merged)
    ordered = sorted(blocks, key=lambda block: getattr(block, "display_order", 0))
    return [(render_text(getattr(block, "content", ""), merged), getattr(block, "block_type", "paragraph")) for block in ordered]


def render_text(content: str, values: dict[str, Any]) -> str:
    def replace(match: re.Match[str]) -> str:
        key = match.group(1)
        return _format_value(values.get(key))

    return _VAR_PATTERN.sub(replace, content or "")


def extract_variables_from_text(text: str) -> list[str]:
    return sorted({match.group(1) for match in _VAR_PATTERN.finditer(text or "")})


def extract_variables_from_docx(content: bytes) -> list[str]:
    from docx import Document

    doc = Document(BytesIO(content))
    found: set[str] = set()
    for text in _iter_doc_text(doc):
        found.update(extract_variables_from_text(text))
    return sorted(found)


def render_docx_template(
    content: bytes,
    values: dict[str, Any],
    variables: Iterable[Any] = (),
    business_type: str = "",
) -> bytes:
    """渲染上传的 DOCX 模板：替换 {{变量}}，保留模板版式。

    解除协议会限制过大的倍数行距，避免 LibreOffice 打印时溢出到第 3 页。
    """
    from docx import Document
    from docx.shared import Mm

    merged = _values_with_defaults(values, variables)
    if business_type:
        merged = enrich_values(business_type, merged)
    doc = Document(BytesIO(content))
    if business_type == "agreement":
        for section in doc.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
    for paragraph in _iter_paragraphs(doc):
        _replace_in_paragraph(paragraph, merged)
        if business_type == "agreement":
            _cap_line_spacing(paragraph)
    for table in _iter_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, merged)
                    if business_type == "agreement":
                        _cap_line_spacing(paragraph)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_docx_plain_text(
    content: bytes,
    values: dict[str, Any],
    variables: Iterable[Any] = (),
    business_type: str = "",
) -> str:
    from docx import Document

    merged = _values_with_defaults(values, variables)
    if business_type:
        merged = enrich_values(business_type, merged)
    doc = Document(BytesIO(content))
    return "\n".join(render_text(text, merged) for text in _iter_doc_text(doc) if text.strip())


def render_docx_preview_blocks(
    content: bytes,
    values: dict[str, Any],
    variables: Iterable[Any] = (),
    business_type: str = "",
) -> list[tuple[str, str]]:
    """Render a stored docx template into stable blocks for HTML preview."""
    from docx import Document

    merged = _values_with_defaults(values, variables)
    if business_type:
        merged = enrich_values(business_type, merged)

    doc = Document(BytesIO(content))
    blocks: list[tuple[str, str]] = []
    header_seen = False

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            text = render_text(paragraph.text, merged).strip()
            if text:
                kind = _docx_preview_block_type(text, -1, business_type)
                blocks.append((text, "title" if kind == "title" else "header"))
                header_seen = kind != "title"
                break
        if header_seen:
            break

    body_index = 0
    for paragraph in doc.paragraphs:
        text = render_text(paragraph.text, merged).strip()
        if not text:
            continue
        kind = _docx_preview_block_type(text, body_index, business_type)
        if kind == "header":
            if not header_seen:
                blocks.append((text, kind))
                header_seen = True
            continue
        blocks.append((text, kind))
        body_index += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = render_text(paragraph.text, merged).strip()
                    if text:
                        kind = _docx_preview_block_type(text, body_index, business_type)
                        if kind == "header":
                            if not header_seen:
                                blocks.append((text, kind))
                                header_seen = True
                            continue
                        blocks.append((text, kind))
                        body_index += 1

    return blocks


def extract_docx_template_blocks(content: bytes, business_type: str = "") -> list[tuple[str, str]]:
    variables = extract_variables_from_docx(content)
    placeholder_values = {code: f"{{{{{code}}}}}" for code in variables}
    return render_docx_preview_blocks(content, placeholder_values, (), business_type)


class _PreviewHtmlParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[tuple[str, str]] = []
        self._current_tag: str | None = None
        self._current_kind: str | None = None
        self._parts: list[str] = []
        self._loose_parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag == "br":
            self._append("\n")
            return
        if self._current_tag is not None:
            return
        kind = self._kind_for(tag, attrs)
        if kind:
            self._current_tag = tag
            self._current_kind = kind
            self._parts = []

    def handle_endtag(self, tag: str) -> None:
        if self._current_tag and tag == self._current_tag:
            self._finish_current()

    def handle_data(self, data: str) -> None:
        self._append(data)

    def close(self) -> None:
        super().close()
        if self._current_tag:
            self._finish_current()
        self._finish_loose()

    def _append(self, data: str) -> None:
        if self._current_tag:
            self._parts.append(data)
        else:
            self._loose_parts.append(data)

    def _finish_current(self) -> None:
        text = _clean_preview_text("".join(self._parts))
        if text:
            if self._current_kind == "pre":
                for line in text.splitlines():
                    line_text = _clean_preview_text(line)
                    if line_text:
                        self.blocks.append((line_text, "paragraph"))
            else:
                self.blocks.append((text, self._current_kind or "paragraph"))
        self._current_tag = None
        self._current_kind = None
        self._parts = []

    def _finish_loose(self) -> None:
        text = _clean_preview_text("".join(self._loose_parts))
        if text:
            for line in text.splitlines():
                line_text = _clean_preview_text(line)
                if line_text:
                    self.blocks.append((line_text, "paragraph"))
        self._loose_parts = []

    @staticmethod
    def _kind_for(tag: str, attrs: list[tuple[str, str | None]]) -> str | None:
        classes = set()
        for key, value in attrs:
            if key == "class" and value:
                classes.update(value.split())
        if tag in {"h1", "h2", "h3"}:
            return "title"
        if tag == "pre" or "template-docx-preview" in classes:
            return "pre"
        if "agr-header" in classes:
            return "header"
        if tag == "p":
            if "agr-head" in classes:
                return "head"
            if {"agr-sign", "cert-sign"} & classes:
                return "sign"
            if {"agr-line", "cert-line"} & classes:
                return "line"
            return "paragraph"
        return None


def _clean_preview_text(text: str) -> str:
    return re.sub(r"[ \t\r\f\v]+", " ", (text or "").replace("\xa0", " ")).strip()


def extract_blocks_from_preview_html(html: str) -> list[tuple[str, str]]:
    parser = _PreviewHtmlParser()
    parser.feed(html or "")
    parser.close()
    return parser.blocks


def render_preview_html_docx(html: str, business_type: str = "") -> bytes:
    """Render a one-off edited preview into docx without changing the stored template."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt

    blocks = extract_blocks_from_preview_html(html)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(31.75)
    section.right_margin = Mm(31.75)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    font_size = Pt(14 if business_type == "income_certificate" else 12)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = font_size
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for text, kind in blocks:
        if kind == "header":
            p = doc.sections[0].header.paragraphs[0]
            p.text = text
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.name = "宋体"
                run.font.size = Pt(9)
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            continue

        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "宋体"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if kind == "title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.bold = True
            run.font.size = Pt(18 if business_type == "income_certificate" else 16)
        else:
            run.font.size = font_size
            p.paragraph_format.line_spacing = 1.45 if business_type == "agreement" else 1.5
            p.paragraph_format.space_after = Pt(4 if business_type == "agreement" else 6)
            if kind in {"body", "paragraph"}:
                p.paragraph_format.first_line_indent = Pt(24)
            if kind == "sign" and business_type == "income_certificate":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if kind == "sign" and business_type == "agreement":
                p.paragraph_format.space_before = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _iter_doc_text(doc: Any) -> Iterable[str]:
    for paragraph in _iter_paragraphs(doc):
        yield paragraph.text
    for table in _iter_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph.text


def _iter_paragraphs(doc: Any) -> Iterable[Any]:
    yield from doc.paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _cap_line_spacing(paragraph: Any, max_ratio: float = 1.3) -> None:
    """压缩过大的倍数行距，缓解 LibreOffice 比 Word 行高偏大导致的溢页。

    仅处理「按倍数」的行距(line_spacing 为浮点倍数);固定值行距(Pt/Mm)和段距不动。
    """
    pf = paragraph.paragraph_format
    ls = pf.line_spacing
    if isinstance(ls, float) and ls > max_ratio:
        pf.line_spacing = max_ratio


def _iter_tables(doc: Any) -> Iterable[Any]:
    yield from doc.tables
    for section in doc.sections:
        yield from section.header.tables
        yield from section.footer.tables


def _docx_preview_block_type(text: str, body_index: int, business_type: str) -> str:
    compact = re.sub(r"\s+", "", text or "")
    if business_type == "agreement":
        if "商业秘密" in compact and "禁止外传" in compact:
            return "header"
        if compact == "解除劳动合同协议书":
            return "title"
        if text.startswith(("甲方（公司）：", "乙方（员工）：")):
            return "head"
        if text.startswith("甲方：") or re.match(r"^年\s+月\s+日", text):
            return "sign"
        if body_index == 0:
            return "title"
        return "paragraph"
    if business_type == "income_certificate":
        if compact == "收入证明":
            return "title"
        if text.startswith(("单位联系电话：", "单位地址：")):
            return "line"
        if body_index == 0:
            return "title"
        if body_index >= 3:
            return "sign"
        return "paragraph"
    if body_index == 0:
        return "title"
    return "paragraph"


def _replace_in_paragraph(paragraph: Any, values: dict[str, Any]) -> None:
    original = paragraph.text
    if "{{" not in original:
        return
    matches = list(_VAR_PATTERN.finditer(original))
    if not matches:
        return
    replacements = [
        (match.start(), match.end(), _format_value(values.get(match.group(1))))
        for match in matches
    ]
    if all(original[start:end] == replacement for start, end, replacement in replacements):
        return
    if not paragraph.runs:
        paragraph.add_run(_replace_text_ranges(original, 0, len(original), replacements))
        return

    cursor = 0
    segments: list[tuple[Any, int, int]] = []
    if paragraph.runs:
        for run in paragraph.runs:
            text = run.text or ""
            start = cursor
            cursor += len(text)
            segments.append((run, start, cursor))

    # If python-docx exposes text not present in runs (for example uncommon inline
    # elements), fall back to the old whole-paragraph replacement rather than
    # risking a partial write.
    if cursor != len(original):
        replaced = render_text(original, values)
        first = paragraph.runs[0]
        first.text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
        return

    for run, start, end in segments:
        run.text = _replace_text_ranges(original, start, end, replacements)


def _format_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (date, datetime)):
        return date_text(value)
    if isinstance(value, Decimal):
        return money_text(value)
    return str(value)


def _replace_text_ranges(
    original: str,
    start: int,
    end: int,
    replacements: list[tuple[int, int, str]],
) -> str:
    cursor = start
    parts: list[str] = []
    for replace_start, replace_end, replacement in replacements:
        if replace_end <= start:
            continue
        if replace_start >= end:
            break
        overlap_start = max(replace_start, start)
        overlap_end = min(replace_end, end)
        if cursor < overlap_start:
            parts.append(original[cursor:overlap_start])
        if start <= replace_start < end:
            parts.append(replacement)
        cursor = max(cursor, overlap_end)
    if cursor < end:
        parts.append(original[cursor:end])
    return "".join(parts)


def _values_with_defaults(values: dict[str, Any], variables: Iterable[Any]) -> dict[str, Any]:
    merged: dict[str, Any] = {}
    for var in variables:
        code = getattr(var, "variable_code", "")
        default = getattr(var, "default_value", None)
        if code and default not in (None, ""):
            merged[code] = default
    merged.update(values)
    return merged


def sample_values(business_type: str) -> dict[str, Any]:
    if business_type == "agreement":
        return enrich_values(
            "agreement",
            {
                "company": "示例公司",
                "name": "张三",
                "id_card": "440300199001010000",
                "dissolve_date": date.today(),
                "last_work_date": date.today(),
                "salary_until": date.today(),
                "base_amount": Decimal("20000"),
                "total_amount": Decimal("80000"),
                "installments": [
                    {"pay_date": date.today(), "amount": Decimal("40000")},
                    {"pay_date": date.today(), "amount": Decimal("40000")},
                ],
            },
        )
    if business_type == "income_certificate":
        return enrich_values(
            "income_certificate",
            {
                "company": "示例公司",
                "name": "张三",
                "id_card": "440300199001010000",
                "position": "高级专员",
                "hire_date": date.today(),
                "leave_date": None,
                "basic_salary": Decimal("20000"),
                "target_bonus": Decimal("60000"),
                "annual_package": Decimal("300000"),
                "issue_date": date.today(),
            },
        )
    return {}


DEFAULT_TEMPLATES = [
    {
        "code": "agreement_release",
        "name": "解除劳动合同协议书",
        "business_type": "agreement",
        "description": "补偿金计算后生成解除劳动合同协议书的默认模板。",
        "version": "1.0",
        "blocks": [
            {"block_type": "header", "content": "商业秘密，禁止外传", "display_order": 10},
            {"block_type": "title", "content": "解除劳动合同协议书", "display_order": 20},
            {"block_type": "head", "content": "甲方（公司）：{{company}}", "display_order": 30},
            {"block_type": "head", "content": "乙方（员工）：{{name}}        身份证号码：{{id_card}}", "display_order": 40},
            {
                "block_type": "paragraph",
                "content": "鉴于甲乙双方签订了《劳动合同》，现经双方友好协商于公司实际经营地签订本协议：",
                "display_order": 50,
            },
            {
                "block_type": "paragraph",
                "content": "1、双方于{{dissolve_date_text}}解除劳动关系，乙方最后工作日为{{last_work_date_text}}；甲方为乙方缴纳的社会保险的最后月份为{{social_security_month}}。",
                "display_order": 60,
            },
            {
                "block_type": "paragraph",
                "content": "2、甲方向乙方支付的工资计算至{{salary_until_text}}，以{{base_amount_text}}元为基数根据该月实际出勤天数据实折算。因解除劳动合同，甲方向乙方支付人民币{{total_amount_text}}元，作为甲方终结与乙方基于劳动关系的全部权利义务的补偿。前述款项支付时间为：{{installments_text}}",
                "display_order": 70,
            },
            {
                "block_type": "paragraph",
                "content": "3、甲方有权从乙方应发工资、补偿金等款项中扣除依法应由乙方承担或应向甲方返还的费用，包括但不限于社保、公积金、个人所得税、缺勤扣款、资产损坏或遗失赔偿、借款及其他应扣款事项。",
                "display_order": 80,
            },
            {
                "block_type": "paragraph",
                "content": "4、乙方应当于最后工作日之前发起离职流程，并逐项办理工作交接、资产归还及相关手续。离职手续应根据公司相关流程办理，经甲方确认后方为办理完成。",
                "display_order": 90,
            },
            {
                "block_type": "paragraph",
                "content": "5、乙方在协商离职、办理离职手续及离职后，仍应遵守劳动合同及保密协议相关规定，不得侵害甲方商业秘密或其他合法权益。",
                "display_order": 100,
            },
            {
                "block_type": "paragraph",
                "content": "特别声明：本协议所涉及的离职事项、工资、补偿金标准及金额等均为涉密信息。乙方不得通过网络媒体或其他渠道对外或对内发表、传播、复制影响公司秩序或损害公司形象的信息。",
                "display_order": 110,
            },
            {
                "block_type": "paragraph",
                "content": "6、如乙方未履行本协议项下任何合同义务，甲方有权依法主张相应权利；如给公司造成损失的，乙方应予以赔偿。",
                "display_order": 120,
            },
            {
                "block_type": "paragraph",
                "content": "7、双方解除劳动关系后，乙方无需承担竞业限制义务，公司无需向乙方支付任何竞业限制补偿金。",
                "display_order": 130,
            },
            {
                "block_type": "paragraph",
                "content": "8、乙方离职时应向甲方真实、准确、完整披露其在职期间表现及工作情况；如乙方隐瞒严重违反规章制度或劳动合同行为，甲方有权解除本协议。",
                "display_order": 140,
            },
            {
                "block_type": "paragraph",
                "content": "9、本协议的成立、效力、解释、执行和争议解决，均适用中国法律法规。本协议为双方解除劳动合同关系及相关事宜的最终真实意思表示，取代此前双方就同一事项的所有协商、协议或备忘录。本协议壹式贰份，双方各执壹份，自公司盖章和员工签字后生效。",
                "display_order": 150,
            },
            {"block_type": "sign", "content": "甲方：{{company}}                乙方：", "display_order": 160},
            {"block_type": "sign", "content": "年     月     日                         年     月     日", "display_order": 170},
        ],
        "variables": [
            {"variable_code": "company", "variable_name": "公司", "source_type": "employee_field", "source_key": "company_name", "required": True},
            {"variable_code": "name", "variable_name": "员工姓名", "source_type": "employee_field", "source_key": "name", "required": True},
            {"variable_code": "id_card", "variable_name": "证件号码", "source_type": "employee_field", "source_key": "id_number"},
            {"variable_code": "dissolve_date_text", "variable_name": "解除日期文本", "source_type": "computed"},
            {"variable_code": "last_work_date_text", "variable_name": "最后工作日文本", "source_type": "computed"},
            {"variable_code": "salary_until_text", "variable_name": "工资截止日文本", "source_type": "computed"},
            {"variable_code": "base_amount_text", "variable_name": "补偿基数文本", "source_type": "computed"},
            {"variable_code": "total_amount_text", "variable_name": "补偿总额文本", "source_type": "computed"},
            {"variable_code": "installments_text", "variable_name": "分期付款文本", "source_type": "computed"},
        ],
    },
    {
        "code": "annual_income",
        "name": "年包收入证明",
        "business_type": "income_certificate",
        "description": "证明开具工具默认收入证明模板。",
        "version": "1.0",
        "blocks": [
            {"block_type": "title", "content": "收入证明", "display_order": 10},
            {
                "block_type": "paragraph",
                "content": "兹证明公司员工{{name}}，身份证号{{id_card}}，{{period_text}}在我公司任职{{position}}，年薪预算总包{{annual_package_wan}}万元（税前），包括：月基本工资税前{{basic_salary_text}}元，如完成绩效目标年终奖金{{target_bonus_text}}元。",
                "display_order": 20,
            },
            {
                "block_type": "paragraph",
                "content": "以上情况属实，公司予以证明。本证明仅用于证明离职员工在我司的工作及在我司的收入基本情况（收入与员工在各考核周期完成绩效目标情况挂钩），不作为我司对员工任何形式的担保文件。",
                "display_order": 30,
            },
            {"block_type": "paragraph", "content": "特此证明！", "display_order": 40},
            {"block_type": "sign", "content": "{{company}}", "display_order": 50},
            {"block_type": "sign", "content": "{{issue_date_text}}", "display_order": 60},
            {"block_type": "line", "content": "单位联系电话：{{company_phone}}", "display_order": 70},
            {"block_type": "line", "content": "单位地址：{{company_address}}", "display_order": 80},
        ],
        "variables": [
            {"variable_code": "company", "variable_name": "公司", "source_type": "employee_field", "source_key": "company_name", "required": True},
            {"variable_code": "name", "variable_name": "员工姓名", "source_type": "employee_field", "source_key": "name", "required": True},
            {"variable_code": "id_card", "variable_name": "证件号码", "source_type": "employee_field", "source_key": "id_number"},
            {"variable_code": "position", "variable_name": "职位", "source_type": "employee_field", "source_key": "position"},
            {"variable_code": "period_text", "variable_name": "任职期间文本", "source_type": "computed"},
            {"variable_code": "annual_package_wan", "variable_name": "年包万元文本", "source_type": "computed"},
            {"variable_code": "basic_salary_text", "variable_name": "月基本工资文本", "source_type": "computed"},
            {"variable_code": "target_bonus_text", "variable_name": "目标年终奖文本", "source_type": "computed"},
            {"variable_code": "issue_date_text", "variable_name": "开具日期文本", "source_type": "computed"},
            {"variable_code": "company_phone", "variable_name": "单位联系电话", "source_type": "fixed", "default_value": COMPANY_PHONE},
            {"variable_code": "company_address", "variable_name": "单位地址", "source_type": "fixed", "default_value": COMPANY_ADDRESS},
        ],
    },
]
