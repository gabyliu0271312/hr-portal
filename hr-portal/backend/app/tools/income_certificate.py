"""收入证明 — 渲染服务

HTML 预览与 docx 下载同源，按「年包收入证明」模板生成。
"""
from __future__ import annotations

from datetime import date
from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO
import base64
from pathlib import Path


TITLE = "收入证明"
COMPANY_PHONE = "0755-86685111"
COMPANY_ADDRESS = "广东省深圳市南山区科苑路15号科兴科学园A3栋16楼"
LOGO_PATH = Path(__file__).with_name("income_certificate_logo.png")
LOGO_WIDTH_INCH = 1.45


def _logo_bytes() -> bytes | None:
    try:
        return LOGO_PATH.read_bytes()
    except Exception:
        return None


def _logo_data_url() -> str:
    data = _logo_bytes()
    if not data:
        return ""
    return "data:image/png;base64," + base64.b64encode(data).decode("ascii")


def _money(v: Decimal) -> str:
    return f"{v.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP):,.2f}"


def _wan(v: Decimal) -> str:
    return f"{(v / Decimal('10000')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP):,.2f}"


def _date_text(d: date | None) -> str:
    if not d:
        return ""
    return f"{d.year}年{d.month}月{d.day}日"


def _period_text(start: date | None, end: date | None) -> str:
    if start and end:
        return f"{_date_text(start)}至{_date_text(end)}"
    if start:
        return f"{_date_text(start)}至今"
    return "至今"


def build_blocks(data: dict) -> list[tuple[str, str]]:
    """返回收入证明正文区块 (文本, 类型)。

    类型 body = 正文缩进段落；line = 单行信息；sign = 落款。
    """
    name = data.get("name") or ""
    id_card = data.get("id_card") or ""
    company = data.get("company") or ""
    position = data.get("position") or ""
    hire_date = data.get("hire_date")
    leave_date = data.get("leave_date")
    basic_salary = Decimal(str(data.get("basic_salary") or 0))
    target_bonus = Decimal(str(data.get("target_bonus") or 0))
    annual_package = Decimal(str(data.get("annual_package") or 0))
    issue_date = data.get("issue_date")

    first = (
        f"兹证明公司员工{name}，身份证号{id_card}，{_period_text(hire_date, leave_date)}"
        f"在我公司任职{position}，年薪预算总包{_wan(annual_package)}万元（税前），包括："
        f"月基本工资税前{_money(basic_salary)}元，如完成绩效目标年终奖金{_money(target_bonus)}元。"
    )
    second = (
        "以上情况属实，公司予以证明。本证明仅用于证明离职员工在我司的工作及在我司的收入基本情况"
        "（收入与员工在各考核周期完成绩效目标情况挂钩），不作为我司对员工任何形式的担保文件。"
    )
    return [
        (first, "body"),
        (second, "body"),
        ("特此证明！", "body"),
        (company, "sign"),
        (_date_text(issue_date), "sign"),
        (f"单位联系电话：{COMPANY_PHONE}", "line"),
        (f"单位地址：{COMPANY_ADDRESS}", "line"),
    ]


def _esc(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace(" ", "&nbsp;")


def _split_blocks(blocks: list[tuple[str, str]]) -> tuple[str, list[tuple[str, str]]]:
    title = TITLE
    body_blocks: list[tuple[str, str]] = []
    for text, kind in blocks:
        if kind == "title":
            title = text
        else:
            body_blocks.append((text, kind))
    return title, body_blocks


def render_html(data: dict, template_blocks: list[tuple[str, str]] | None = None) -> str:
    title, blocks = _split_blocks(template_blocks or build_blocks(data))
    parts = []
    for text, kind in blocks:
        cls = {"body": "cert-p", "paragraph": "cert-p", "sign": "cert-sign", "line": "cert-line", "footer": "cert-line"}.get(kind, "cert-p")
        parts.append(f'<p class="{cls}">{_esc(text)}</p>')
    body = "".join(parts)
    logo = _logo_data_url()
    logo_html = f'<div class="cert-header"><img class="cert-logo" src="{logo}" /></div>' if logo else ""
    return f"""<div class="cert-doc">
{logo_html}
<h1 class="cert-title">{_esc(title)}</h1>
{body}
</div>"""


def render_docx(data: dict, template_blocks: list[tuple[str, str]] | None = None) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt

    title_text, blocks = _split_blocks(template_blocks or build_blocks(data))
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    logo = _logo_bytes()
    if logo:
        header_p = section.header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header_p.add_run()
        run.add_picture(BytesIO(logo), width=Inches(LOGO_WIDTH_INCH))

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for text, kind in blocks:
        p = doc.add_paragraph(text)
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        if kind in {"body", "paragraph"}:
            pf.first_line_indent = Pt(24)
        elif kind == "sign":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
