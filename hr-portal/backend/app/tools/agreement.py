"""解除劳动合同协议 — 渲染服务

协议正文与排版只在此处维护一份，HTML 预览与 docx 下载同源，保证两者格式一致。
"""
from __future__ import annotations

import calendar
from datetime import date
from decimal import Decimal, ROUND_DOWN
from io import BytesIO


HEADER_TEXT = "商业秘密，禁止外传"
TITLE = "解除劳动合同协议书"


def _add_months(d: date, months: int) -> date:
    idx = d.month - 1 + months
    year = d.year + idx // 12
    month = idx % 12 + 1
    day = min(d.day, calendar.monthrange(year, month)[1])
    return date(year, month, day)


def _money(v: Decimal) -> str:
    return f"{v.quantize(Decimal('0.01'), rounding=ROUND_DOWN):,.2f}"


def compute_installments(
    total: Decimal,
    rules: list[dict],
    leave_date: date,
) -> list[dict]:
    """按规则计算分期：前面各期向下取整，最后一期取剩余，保证合计=total。

    rules: [{"period_no","ratio","months_after","pay_day"}, ...] 已按 period_no 排序
    返回: [{"pay_date": date, "amount": Decimal}, ...]
    """
    if not rules:
        return [{"pay_date": _add_months(leave_date, 1).replace(day=min(15, 28)), "amount": total}]
    result: list[dict] = []
    allocated = Decimal("0")
    n = len(rules)
    for i, r in enumerate(rules):
        if i == n - 1:
            amount = total - allocated
        else:
            amount = (total * Decimal(str(r["ratio"])) / Decimal("100")).quantize(
                Decimal("1"), rounding=ROUND_DOWN
            )
            allocated += amount
        months = int(r["months_after"])
        pay_day = int(r["pay_day"])
        base = _add_months(leave_date, months)
        last_day = calendar.monthrange(base.year, base.month)[1]
        pay_date = base.replace(day=min(pay_day, last_day))
        result.append({"pay_date": pay_date, "amount": amount})
    return result


def _installment_sentence(installments: list[dict]) -> str:
    parts = []
    for it in installments:
        d = it["pay_date"]
        parts.append(f"{d.year}年{d.month}月{d.day}日前支付{_money(it['amount'])}元")
    return "；".join(parts) + "。"


def build_blocks(data: dict) -> list[tuple[str, str]]:
    """返回协议正文区块列表 (文本, 类型)。

    类型 head = 抬头行（顶格不缩进、整行不折行）；body = 正文段落（首行缩进）。
    data 需含：company, name, id_card, dissolve_date(date), last_work_date(date),
    social_security_month(str), salary_until(date), base_amount(Decimal),
    total_amount(Decimal), installments(list)
    """
    dd = data["dissolve_date"]
    lw = data["last_work_date"]
    su = data["salary_until"]
    base = data["base_amount"]
    total = data["total_amount"]
    inst = _installment_sentence(data["installments"])

    head = [
        (f"甲方（公司）：{data['company']}", "head"),
        (f"乙方（员工）：{data['name']}        身份证号码：{data['id_card']}", "head"),
    ]
    body = [
        "鉴于甲乙双方签订了《劳动合同》，现经双方友好协商于公司实际经营地签订本协议：",
        f"1、双方于{dd.year}年{dd.month}月{dd.day}日解除劳动关系，乙方最后工作日为"
        f"{lw.year}年{lw.month}月{lw.day}日；甲方为乙方缴纳的社会保险的最后月份为"
        f"{data['social_security_month']}。",
        f"2、甲方向乙方支付的工资计算至{su.year}年{su.month}月{su.day}日，以{_money(base)}元为基数"
        f"根据该月实际出勤天数据实折算。因解除劳动合同，甲方向乙方支付人民币{_money(total)}元，"
        "作为甲方终结与乙方基于劳动关系的全部权利义务的补偿，乙方确认该款项已经包括但不限于下列内容："
        "在职期间甲方应付未付给的工资、经济补偿金、应缴未缴的住房公积金、未休年假工资，未休假补贴、"
        "加班费、报销款、奖金提成、补贴津贴、医疗补偿金、股票期权等全部费用。除此外，甲方或甲方关联公司"
        "无需向乙方支付任何其他费用，乙方免除甲方或甲方关联公司其他全部债务，如乙方曾在甲方关联公司任职的，"
        "乙方不再向甲方关联公司主张任何权利。前述款项支付时间为：" + inst,
        "3、甲方有权从乙方应发工资、补偿金等扣除下列费用。扣除款项为：（1）社保、公积金扣款，"
        "包括公司代员工垫付部分；（2）个人所得税，包括公司代员工垫付部分；（3）缺勤、旷工、怠工、"
        "迟到早退等扣款；（4）公司资产遗失或损坏扣款；（5）企业打车扣款及差旅超标扣款；（6）给公司"
        "造成损失后应向公司支付的赔偿款；（7）员工借款；（8）工卡补办费用；（9）其他发生的扣款事项。",
        "4、乙方应当于最后工作日之前发起离职流程并逐项办理工作交接和资产归还等手续。其中工作交接"
        "包括但不限于：向公司归还所有与工作有关的记录、文案、账户、密码、图册以及工作日志、项目报告、"
        "客户名单等所有纸质和电子文件等资料；资产归还包括但不限于向公司归还办公设备、钥匙、出入证及"
        "还清借支资金等。离职手续应根据公司相关流程和手续进行，并经甲方确认后方为办理完成；办理离职"
        "手续期间不视为乙方工作时间。",
        "5、乙方在协商离职、办理离职手续及离职后等情况，仍应遵守劳动合同及其《保密协议》相关规定，"
        "严禁采取复制、下载、破坏、泄露、反向工程、使用等手段侵害甲方涉密信息，亦不得做出有损于甲方"
        "合法权益的行为。",
        "特别声明，本协议所涉及的有关离职全部事项均为涉密信息，包括但不限于工资、补偿金标准、"
        "补偿金金额等。",
        "乙方不得通过网络媒介或其他媒介对外或对内发表、传播、复制影响公司团结，损害公司形象，"
        "揭露个人隐私，阻碍人员调整、否定高管人员等言论；或组织煽动员工采取非法律手段解决争议纠纷或"
        "聚集维权等。",
        "6、如乙方未履行本协议项下的任何合同义务（包括但不限于保密义务），则甲方有权没收本协议项下"
        "所有应付款。如给公司造成损失的，员工应予以赔偿。",
        "7、双方解除劳动关系后，乙方无需承担竞业限制义务，公司无需向员工支付任何竞业限制补偿金。",
        "8、乙方离职时应向甲方真实、准确、完整披露其在职期间表现及工作，如乙方离职时隐瞒在职期间"
        "存在严重违反规章制度或劳动合同行为的，则甲方有权解除本协议。",
        "9、本协议的成立、效力、解释、执行和争议解决，均应适用中国法律法规并受本协议签订地法院管辖。"
        "本协议为双方解除劳动合同关系及相关事宜的最终真实意思表示，取代之前双方就同一事项的所有书面或"
        "口头的协商、协议、备忘录等。本协议壹式贰份，双方各执壹份，自公司盖章和员工签字后生效。",
    ]
    return head + [(t, "body") for t in body]


def _esc(s: str) -> str:
    return (
        s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace(" ", "&nbsp;")
    )


def _split_blocks(blocks: list[tuple[str, str]]) -> tuple[str, str, list[tuple[str, str]]]:
    header = HEADER_TEXT
    title = TITLE
    body_blocks: list[tuple[str, str]] = []
    for text, kind in blocks:
        if kind == "header":
            header = text
        elif kind == "title":
            title = text
        else:
            body_blocks.append((text, kind))
    return header, title, body_blocks


def _default_blocks(data: dict) -> list[tuple[str, str]]:
    return build_blocks(data) + [
        (f"甲方：{data.get('company') or ''}                乙方：", "sign"),
        ("年     月     日                         年     月     日", "sign"),
    ]


def render_html(data: dict, template_blocks: list[tuple[str, str]] | None = None) -> str:
    """与 docx 同源的 A4 预览 HTML（含打印样式，宋体）。"""
    header, title, blocks = _split_blocks(template_blocks or _default_blocks(data))
    parts = []
    for text, kind in blocks:
        cls = (
            "agr-p"
            if kind in {"body", "paragraph"}
            else "agr-head"
            if kind == "head"
            else "agr-sign"
            if kind == "sign"
            else "agr-line"
        )
        parts.append(f'<p class="{cls}">{_esc(text)}</p>')
    body = "".join(parts)
    return f"""<div class="agr-doc">
<div class="agr-header">{_esc(header)}</div>
<h1 class="agr-title">{_esc(title)}</h1>
{body}
</div>"""


def render_docx(data: dict, template_blocks: list[tuple[str, str]] | None = None) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt

    header, title_text, blocks = _split_blocks(template_blocks or _default_blocks(data))
    doc = Document()
    # 页眉：小五(9pt) 宋体 靠右
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(31.75)
    section.right_margin = Mm(31.75)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    header_p = section.header.paragraphs[0]
    header_p.text = header
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in header_p.runs:
        r.font.size = Pt(9)
        r.font.name = "宋体"
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 正文默认字体：宋体 小四(12pt)，中英文统一
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "宋体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for text, kind in blocks:
        p = doc.add_paragraph(text)
        pf = p.paragraph_format
        pf.line_spacing = 1.45
        pf.space_after = Pt(4)
        if kind in {"body", "paragraph"}:
            pf.first_line_indent = Pt(24)  # 首行缩进 2 字符（小四 12pt）
        elif kind == "head":
            pf.space_after = Pt(3)
        elif kind == "sign":
            pf.space_before = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
