from datetime import date
from io import BytesIO

from docx import Document
from docx.shared import Mm

from app.tools import router as tools_router
from app.tools.models import DocumentTemplate, DocumentTemplateBlock


def _docx_bytes(doc: Document) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_agreement_render_uses_uploaded_docx_layout_when_preview_is_not_edited():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Mm(12)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 2
    paragraph.add_run("Company: {{company}}")
    content = _docx_bytes(doc)
    template = DocumentTemplate(
        code="agreement_release",
        name="Agreement",
        business_type="agreement",
        template_file=content,
        parsed_variables=[],
        layout_config={},
    )
    template.variables = []
    template.blocks = []

    rendered = tools_router._render_agreement_docx_content(
        tools_router.AgreementData(
            template_code="agreement_release",
            template_name="Agreement",
            company="Example Co.",
            name="Alice",
            id_card="ID001",
            dissolve_date=date(2026, 6, 16),
            last_work_date=date(2026, 6, 16),
            social_security_month="2026-06",
            salary_until=date(2026, 6, 16),
            base_amount=20000,
            total_amount=80000,
            installments=[
                tools_router.AgreementInstallment(pay_date=date(2026, 7, 15), amount=80000)
            ],
        ),
        template,
    )
    out = Document(BytesIO(rendered))

    assert abs(out.sections[0].left_margin - Mm(12)) < 500
    assert out.paragraphs[0].paragraph_format.line_spacing == 1.3
    assert out.paragraphs[0].text == "Company: Example Co."


def test_agreement_render_uses_saved_blocks_after_preview_save():
    doc = Document()
    doc.add_paragraph("Company: {{company}}")
    content = _docx_bytes(doc)
    template = DocumentTemplate(
        code="agreement_release",
        name="Agreement",
        business_type="agreement",
        template_file=content,
        parsed_variables=[],
        layout_config={"render_mode": "standard_blocks"},
    )
    template.variables = []
    template.blocks = [
        DocumentTemplateBlock(block_type="title", content="Saved Agreement", display_order=10)
    ]

    rendered = tools_router._render_agreement_docx_content(
        tools_router.AgreementData(
            template_code="agreement_release",
            template_name="Agreement",
            company="Example Co.",
            name="Alice",
            id_card="ID001",
            dissolve_date=date(2026, 6, 16),
            last_work_date=date(2026, 6, 16),
            social_security_month="2026-06",
            salary_until=date(2026, 6, 16),
            base_amount=20000,
            total_amount=80000,
            installments=[],
        ),
        template,
    )
    out = Document(BytesIO(rendered))

    assert out.paragraphs[0].text == "Saved Agreement"
