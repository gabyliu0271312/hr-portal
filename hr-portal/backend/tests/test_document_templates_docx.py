from datetime import date
from io import BytesIO

from docx import Document

from app.tools import router as tools_router
from app.tools.models import DocumentTemplate
from app.tools.document_templates import extract_variables_from_docx, render_docx_template


def _docx_bytes(doc: Document) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_render_docx_template_replaces_split_run_variables_without_flattening_paragraph():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("乙方（员工）：")
    p.add_run("{{name}}        ")
    p.add_run("身份证号码：")
    p.add_run("{{")
    p.add_run("id_card")
    p.add_run("}}")
    before_run_count = len(p.runs)

    content = _docx_bytes(doc)

    assert extract_variables_from_docx(content) == ["id_card", "name"]

    rendered = render_docx_template(
        content,
        {"name": "张三", "id_card": "440300199001010000"},
        business_type="agreement",
    )
    out = Document(BytesIO(rendered))
    paragraph = out.paragraphs[0]

    assert paragraph.text == "乙方（员工）：张三        身份证号码：440300199001010000"
    assert len(paragraph.runs) == before_run_count
    assert paragraph.runs[0].text == "乙方（员工）："
    assert paragraph.runs[1].text == "张三        "
    assert paragraph.runs[2].text == "身份证号码："
    assert paragraph.runs[3].text == "440300199001010000"
    assert paragraph.runs[4].text == ""
    assert paragraph.runs[5].text == ""


def test_render_docx_template_replaces_placeholder_split_before_closing_braces():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("双方于")
    p.add_run("{{dissolve_date_text")
    p.add_run("}}")
    p.add_run("解除劳动关系")
    content = _docx_bytes(doc)

    assert extract_variables_from_docx(content) == ["dissolve_date_text"]

    rendered = render_docx_template(
        content,
        {"dissolve_date": date(2026, 6, 16)},
        business_type="agreement",
    )
    out = Document(BytesIO(rendered))
    paragraph = out.paragraphs[0]

    assert paragraph.text == "双方于2026年6月16日解除劳动关系"
    assert [run.text for run in paragraph.runs] == [
        "双方于",
        "2026年6月16日",
        "",
        "解除劳动关系",
    ]


def test_uploaded_agreement_word_preview_uses_agreement_html_layout():
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "商业秘密，禁止外传"
    doc.add_paragraph("解除劳动合同协议书")
    doc.add_paragraph("甲方（公司）：{{company}}")
    doc.add_paragraph("乙方（员工）：{{name}}        身份证号码：{{id_card}}")
    doc.add_paragraph("1、双方于{{dissolve_date_text}}解除劳动关系。")
    content = _docx_bytes(doc)
    template = DocumentTemplate(
        code="agreement_release",
        name="解除劳动合同协议书",
        business_type="agreement",
        template_file=content,
        parsed_variables=[],
        layout_config={},
    )
    template.variables = []

    html = tools_router._render_template_html(
        template,
        {
            "company": "示例公司",
            "name": "张三",
            "id_card": "440300199001010000",
            "dissolve_date": date(2026, 6, 16),
        },
    )

    assert "<pre" not in html
    assert 'class="agr-doc"' in html
    assert 'class="agr-header">商业秘密，禁止外传' in html
    assert "甲方（公司）：示例公司" in html
    assert "1、双方于2026年6月16日解除劳动关系。" in html
